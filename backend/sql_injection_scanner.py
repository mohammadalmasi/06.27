from flask import request, jsonify, send_file
import requests
import re
import os
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from sql_injection_detector import SQLInjectionDetector
from enhanced_sql_injection_detector import EnhancedSQLInjectionDetector
from sonarqube_security_standards import SecurityStandards, SQCategory, VulnerabilityProbability
from datetime import datetime
import json
import tempfile

# Initialize both detectors
ast_detector = SQLInjectionDetector()
enhanced_detector = EnhancedSQLInjectionDetector()

# Based on SonarQube's actual SecurityStandards.java
SQL_INJECTION = ("sql-injection", VulnerabilityProbability.HIGH)
# Maps to CWE-89, CWE-564, CWE-943
# Maps to OWASP A03:2021-Injection

def highlight_sql_injection(code):
    """Highlight vulnerable SQL and NoSQL patterns in code (JS/Python-like)."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def highlight_sql_injection_web(code):
    """Highlight vulnerable SQL and NoSQL patterns in code for web display."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def highlight_sql_injection_word(code):
    """Highlight vulnerable SQL and NoSQL patterns in code for Word document generation."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'[VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE)
    return highlighted

def scan_code_file_ast(filename):
    """Scan code file using AST-based SQL injection detector"""
    try:
        vulnerabilities = ast_detector.scan_file(filename)
        return vulnerabilities
    except Exception as e:
        print(f"AST detection error: {e}")
        return []

def scan_code_file(filename):
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            code = f.read()
    except UnicodeDecodeError:
        with open(filename, 'r', encoding='latin-1') as f:
            code = f.read()
    
    # Get AST-based vulnerabilities
    ast_vulnerabilities = scan_code_file_ast(filename)
    
    # More comprehensive patterns for SQL injection detection (regex-based)
    patterns = [
        # String concatenation with user input
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        
        # F-string formatting
        r'\{\w+\}',     # {user_input} in f-strings
        
        # String formatting
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        
        # Direct SQL construction
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        
        # SQL statements with concatenation
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(DELETE\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',       # DELETE FROM users WHERE id = + user_input
        
        # Framework-specific patterns
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    
    regex_vulnerabilities = []
    lines = code.split('\n')
    
    for line_num, line in enumerate(lines, 1):
        for pattern in patterns:
            matches = re.finditer(pattern, line, re.IGNORECASE)
            for match in matches:
                vuln_info = {
                    'line_number': line_num,
                    'vulnerable_code': match.group(0),
                    'line_content': line.strip(),
                    'pattern': pattern,
                    'type': 'regex_detection'
                }
                regex_vulnerabilities.append(vuln_info)
    
    # Combine AST and regex vulnerabilities
    all_vulnerabilities = []
    
    # Add AST vulnerabilities
    for vuln in ast_vulnerabilities:
        all_vulnerabilities.append({
            'line_number': vuln.get('line_number', 0),
            'vulnerable_code': vuln.get('vulnerable_code', ''),
            'line_content': vuln.get('line_content', ''),
            'description': vuln.get('description', ''),
            'type': 'ast_detection',
            'severity': vuln.get('severity', 'medium')
        })
    
    # Add regex vulnerabilities
    for vuln in regex_vulnerabilities:
        all_vulnerabilities.append({
            'line_number': vuln['line_number'],
            'vulnerable_code': vuln['vulnerable_code'],
            'line_content': vuln['line_content'],
            'description': f"Potential SQL injection vulnerability detected in pattern: {vuln['vulnerable_code']}",
            'type': 'regex_detection',
            'severity': 'medium'
        })
    
    return all_vulnerabilities

def scan_code_content_enhanced(code_content: str, source_name: str) -> dict:
    """Enhanced scanning of code content with SonarQube security standards"""
    try:
        # Create temporary file for scanning
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(code_content)
        temp_file.close()
        
        # Create a fresh detector instance for each scan to avoid cached vulnerabilities
        fresh_detector = EnhancedSQLInjectionDetector()
        vulnerabilities = fresh_detector.scan_file(temp_file.name)
        
        # Clean up temporary file
        os.unlink(temp_file.name)
        
        # Generate enhanced report
        report = fresh_detector.get_enhanced_report()
        summary = report.get('summary', {})
        
        # Generate highlighted code for React frontend
        highlighted_code = None
        original_code = code_content
        file_name = source_name
        
        if vulnerabilities:
            # Use existing highlight function to highlight vulnerable patterns
            highlighted_code = highlight_sql_injection_web(code_content)
            
            # Escape HTML entities in the original code for safe display
            original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Extract filename from source
            if '/' in source_name:
                file_name = source_name.split('/')[-1]
            elif source_name.startswith('http'):
                file_name = source_name.split('/')[-1] if '/' in source_name else 'scanned_code.py'
        
        # Format results with both enhanced and UI-compatible formats
        results = {
            'source': source_name,
            'scan_type': 'enhanced',
            'summary': summary,
            'compliance': report.get('compliance', {}),
            'vulnerabilities': [vuln.to_dict() for vuln in vulnerabilities],
            'total_vulnerabilities': len(vulnerabilities),
            'scan_timestamp': datetime.now().isoformat(),
            
            # Additional UI-compatible fields
            'total_issues': len(vulnerabilities),
            'high_severity': summary.get('high_severity', 0),
            'medium_severity': summary.get('medium_severity', 0),
            'low_severity': summary.get('low_severity', 0),
            'high_count': summary.get('high', 0),
            'medium_count': summary.get('medium', 0),
            'low_count': summary.get('low', 0),
            
            # Code highlighting fields for React frontend
            'highlighted_code': highlighted_code,
            'original_code': original_code,
            'file_name': file_name
        }
        
        return results
        
    except Exception as e:
        return {
            'error': f'Error during enhanced scan: {str(e)}',
            'source': source_name,
            'scan_type': 'enhanced',
            'vulnerabilities': [],
            'total_vulnerabilities': 0,
            'total_issues': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'highlighted_code': None,
            'original_code': '',
            'file_name': source_name
        }

def is_github_py_url(url):
    """Check if URL is a GitHub Python file URL"""
    return 'github.com' in url and url.endswith('.py')

def github_raw_url(url):
    """Convert GitHub blob URL to raw URL"""
    # e.g. https://github.com/user/repo/blob/main/file.py -> https://raw.githubusercontent.com/user/repo/main/file.py
    return url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')

def api_scan_code(current_user):
    """API endpoint for scanning code"""
    try:
        data = request.get_json()
        code_content = data.get('code')
        url = data.get('url')
        
        # Ensure directories exist
        os.makedirs('results', exist_ok=True)
        
        if code_content:
            # Scan provided code content
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
            temp_file.write(code_content)
            temp_file.close()
            
            vulnerabilities = scan_code_file(temp_file.name)
            
            # Clean up temporary file
            os.unlink(temp_file.name)
            
            # Generate highlighted code
            highlighted_code = highlight_sql_injection_web(code_content)
            
            # Escape HTML entities in the original code for safe display
            original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            results = {
                'source': 'Direct input',
                'vulnerabilities': vulnerabilities,
                'total_issues': len(vulnerabilities),
                'high_count': sum(1 for v in vulnerabilities if v.get('severity') == 'high'),
                'medium_count': sum(1 for v in vulnerabilities if v.get('severity') == 'medium'),
                'low_count': sum(1 for v in vulnerabilities if v.get('severity') == 'low'),
                'highlighted_code': highlighted_code,
                'original_code': original_code,
                'file_name': 'scanned_code.py'
            }
            
        elif url:
            # Scan URL content
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        code_content = response.text
                        
                        # Create temporary file for scanning
                        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
                        temp_file.write(code_content)
                        temp_file.close()
                        
                        vulnerabilities = scan_code_file(temp_file.name)
                        
                        # Clean up temporary file
                        os.unlink(temp_file.name)
                        
                        # Generate highlighted code
                        highlighted_code = highlight_sql_injection_web(code_content)
                        
                        # Escape HTML entities in the original code for safe display
                        original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                        
                        # Extract filename from URL
                        file_name = url.split('/')[-1]
                        
                        results = {
                            'source': url,
                            'vulnerabilities': vulnerabilities,
                            'total_issues': len(vulnerabilities),
                            'high_count': sum(1 for v in vulnerabilities if v.get('severity') == 'high'),
                            'medium_count': sum(1 for v in vulnerabilities if v.get('severity') == 'medium'),
                            'low_count': sum(1 for v in vulnerabilities if v.get('severity') == 'low'),
                            'highlighted_code': highlighted_code,
                            'original_code': original_code,
                            'file_name': file_name
                        }
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'No code or URL provided'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during scan: {str(e)}'}), 500

def api_enhanced_scan(current_user):
    """API endpoint for enhanced scanning"""
    try:
        data = request.get_json()
        code_content = data.get('code')
        url = data.get('url')
        
        if code_content:
            # Scan provided code content
            results = scan_code_content_enhanced(code_content, 'Direct input')
            
        elif url:
            # Scan URL content
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        results = scan_code_content_enhanced(response.text, url)
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'Invalid scan parameters'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during enhanced scan: {str(e)}'}), 500

def api_generate_report(current_user):
    """API endpoint for generating Word reports"""
    try:
        data = request.get_json()
        vulnerabilities = data.get('vulnerabilities', [])
        source = data.get('source', 'Unknown')
        
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('SQL Injection Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Source'
        info_table.cell(0, 1).text = source
        info_table.cell(1, 0).text = 'Report Generated'
        info_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        info_table.cell(2, 0).text = 'Total Vulnerabilities'
        info_table.cell(2, 1).text = str(len(vulnerabilities))
        info_table.cell(3, 0).text = 'Risk Level'
        info_table.cell(3, 1).text = 'High' if len(vulnerabilities) > 0 else 'Low'
        
        # Add summary
        doc.add_heading('Executive Summary', level=1)
        summary_text = f"""
This report presents the results of a comprehensive SQL injection vulnerability analysis performed on the provided code. 
The analysis identified {len(vulnerabilities)} potential security vulnerabilities that could allow unauthorized database access 
or manipulation if exploited by malicious actors.

The vulnerabilities identified pose significant security risks and should be addressed immediately to prevent potential 
data breaches, unauthorized access, and system compromise.
        """
        doc.add_paragraph(summary_text.strip())
        
        # Add vulnerability details
        doc.add_heading('Vulnerability Details', level=1)
        
        for i, vuln in enumerate(vulnerabilities, 1):
            doc.add_heading(f'Vulnerability #{i}', level=2)
            
            # Vulnerability info table
            vuln_table = doc.add_table(rows=4, cols=2)
            vuln_table.style = 'Table Grid'
            
            vuln_table.cell(0, 0).text = 'Line Number'
            vuln_table.cell(0, 1).text = str(vuln.get('line_number', 'N/A'))
            vuln_table.cell(1, 0).text = 'Severity'
            vuln_table.cell(1, 1).text = vuln.get('severity', 'Medium').title()
            vuln_table.cell(2, 0).text = 'Detection Method'
            vuln_table.cell(2, 1).text = vuln.get('type', 'Pattern matching').title()
            vuln_table.cell(3, 0).text = 'Description'
            vuln_table.cell(3, 1).text = vuln.get('description', 'SQL injection vulnerability detected')
            
            # Vulnerable code
            doc.add_paragraph('Vulnerable Code:', style='Heading 3')
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(vuln.get('line_content', ''))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            
            # Highlight vulnerable part
            if vuln.get('vulnerable_code'):
                doc.add_paragraph('Vulnerable Pattern:', style='Heading 3')
                vuln_para = doc.add_paragraph()
                vuln_run = vuln_para.add_run(vuln.get('vulnerable_code', ''))
                vuln_run.font.name = 'Courier New'
                vuln_run.font.size = Pt(10)
                vuln_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
            
            doc.add_paragraph()  # Add spacing
        
        # Add recommendations
        doc.add_heading('Recommendations', level=1)
        recommendations = """
1. Use Parameterized Queries: Replace string concatenation with parameterized queries or prepared statements.
2. Input Validation: Implement strict input validation and sanitization for all user inputs.
3. Least Privilege: Ensure database connections use accounts with minimal required permissions.
4. Regular Updates: Keep database software and frameworks updated to the latest security patches.
5. Code Review: Implement mandatory security code reviews for all database interaction code.
6. Security Testing: Conduct regular penetration testing and vulnerability assessments.
        """
        doc.add_paragraph(recommendations.strip())
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sql_injection_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error generating report: {str(e)}'}), 500

def api_sonarqube_export(current_user):
    """API endpoint for SonarQube export"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        vulnerabilities = data.get('vulnerabilities', [])
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Convert to SonarQube format
        sonar_issues = []
        for vuln in vulnerabilities:
            sonar_issue = {
                "engineId": "python-security-scanner",
                "ruleId": vuln.get('rule_key', 'python:S2077'),
                "severity": vuln.get('severity', 'MAJOR'),
                "type": "VULNERABILITY",
                "primaryLocation": {
                    "message": vuln.get('description', 'SQL Injection vulnerability'),
                    "filePath": vuln.get('file_path', 'unknown'),
                    "textRange": {
                        "startLine": vuln.get('line_number', 1),
                        "endLine": vuln.get('line_number', 1)
                    }
                },
                "cwe": vuln.get('cwe_references', []),
                "owasp": vuln.get('owasp_references', []),
                "confidence": vuln.get('confidence', 0.5)
            }
            sonar_issues.append(sonar_issue)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json')
        json.dump({"issues": sonar_issues}, temp_file, indent=2)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sonarqube_issues_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            mimetype='application/json'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error exporting SonarQube format: {str(e)}'}), 500

def api_get_security_standards(current_user):
    """API endpoint for getting security standards"""
    try:
        standards = {
            "sq_categories": [
                {
                    "key": category.key,
                    "name": category.name.replace('_', ' ').title(),
                    "vulnerability_probability": category.vulnerability.name
                }
                for category in SQCategory
            ],
            "vulnerability_probabilities": [
                {
                    "name": prob.name,
                    "score": prob.value
                }
                for prob in VulnerabilityProbability
            ],
            "cwe_mappings": {
                "sql_injection": ["89", "564", "943"],
                "nosql_injection": ["89", "943"],
                "command_injection": ["77", "78", "88", "214"]
            },
            "owasp_top10_2021": [
                "A01:2021-Broken Access Control",
                "A02:2021-Cryptographic Failures",
                "A03:2021-Injection",
                "A04:2021-Insecure Design",
                "A05:2021-Security Misconfiguration",
                "A06:2021-Vulnerable and Outdated Components",
                "A07:2021-Identification and Authentication Failures",
                "A08:2021-Software and Data Integrity Failures",
                "A09:2021-Security Logging and Monitoring Failures",
                "A10:2021-Server-Side Request Forgery"
            ]
        }
        
        return jsonify(standards)
        
    except Exception as e:
        return jsonify({'error': f'Error getting security standards: {str(e)}'}), 500 