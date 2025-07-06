from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import re
import os
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import sqlite3
from sql_injection_detector import SQLInjectionDetector
from enhanced_sql_injection_detector import EnhancedSQLInjectionDetector
from sonarqube_security_standards import SecurityStandards, SQCategory, VulnerabilityProbability
from datetime import datetime, timedelta
import jwt
from functools import wraps
import json
import tempfile
import zipfile

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # 2MB upload limit
app.config['JWT_SECRET_KEY'] = os.environ.get('JWT_SECRET_KEY', 'sql-injection-scanner-secret-key-2024')

# Enable CORS for all origins (production and development)
CORS(app, origins=["*"], 
     methods=["GET", "POST", "OPTIONS"], 
     allow_headers=["Content-Type", "Authorization"])

# Initialize both detectors
ast_detector = SQLInjectionDetector()
enhanced_detector = EnhancedSQLInjectionDetector()

# Authentication configuration
ADMIN_USERNAME = "admin"
ADMIN_PASSWORD = "a"

# Based on SonarQube's actual SecurityStandards.java
SQL_INJECTION = ("sql-injection", VulnerabilityProbability.HIGH)
# Maps to CWE-89, CWE-564, CWE-943
# Maps to OWASP A03:2021-Injection

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        # JWT is passed in the request header
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            try:
                token = auth_header.split(" ")[1]  # Bearer <token>
            except IndexError:
                return jsonify({'error': 'Token is missing!'}), 401
        
        if not token:
            return jsonify({'error': 'Token is missing!'}), 401
        
        try:
            # Decode the token
            data = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
            current_user = data['username']
        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'Token has expired!'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'error': 'Token is invalid!'}), 401
        
        return f(current_user, *args, **kwargs)
    
    return decorated

def highlight_sql_injection(code):
    """Highlight vulnerable SQL and NoSQL patterns in code (JS/Python-like)."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def highlight_sql_injection_web(code):
    """Highlight vulnerable SQL and NoSQL patterns in code for web display."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def highlight_sql_injection_word(code):
    """Highlight vulnerable SQL and NoSQL patterns in code for Word document generation."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'[VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE)
    return highlighted

def ensure_dirs():
    os.makedirs('results', exist_ok=True)
    ensure_dirs()
    # No longer saving to sourcecodes directory
    return None

def scan_code_file_ast(filename):
    """Scan code file using AST-based SQL injection detector"""
    try:
        vulnerabilities = ast_detector.scan_file(filename)
        return vulnerabilities
    except Exception as e:
        print(f"AST detection error: {e}")
        return []

def scan_code_file(filename):
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            code = f.read()
    except UnicodeDecodeError:
        with open(filename, 'r', encoding='latin-1') as f:
            code = f.read()
    
    # Get AST-based vulnerabilities
    ast_vulnerabilities = scan_code_file_ast(filename)
    
    # More comprehensive patterns for SQL injection detection (regex-based)
    patterns = [
        # String concatenation with user input
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        
        # F-string formatting
        r'\{\w+\}',     # {user_input} in f-strings
        
        # String formatting
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        
        # Database query patterns with user input
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        
        # Raw SQL queries with variables
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        
        # Common vulnerable patterns
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
    ]
    
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'[VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE)
    
    # Create comprehensive report with both detection methods
    base = os.path.basename(filename)
    docx_filename = os.path.join('results', base.replace('.', '_result.', 1).rsplit('.', 1)[0] + '.docx')
    doc = Document()
    
    # Add title
    title = doc.add_heading('SQL Injection Analysis Report', 0)
    
    # Add file information
    doc.add_heading('File Information', level=1)
    doc.add_paragraph(f'File: {filename}')
    doc.add_paragraph(f'Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add AST-based vulnerabilities
    if ast_vulnerabilities:
        doc.add_heading('AST-Based Detection Results', level=1)
        doc.add_paragraph(f'Found {len(ast_vulnerabilities)} vulnerabilities using AST analysis:')
        
        for vuln in ast_vulnerabilities:
            # Vulnerability details
            p = doc.add_paragraph()
            p.add_run(f'Line {vuln.line_number}: ').bold = True
            p.add_run(f'{vuln.description}')
            
            # Severity and confidence
            p2 = doc.add_paragraph()
            p2.add_run(f'Severity: {vuln.severity} | Confidence: {vuln.confidence:.2f}')
            
            # Code snippet
            if vuln.code_snippet:
                doc.add_paragraph('Code Snippet:')
                code_para = doc.add_paragraph(vuln.code_snippet)
                code_para.style = 'No Spacing'
            
            # Remediation
            doc.add_paragraph(f'Remediation: {vuln.remediation}')
            doc.add_paragraph('---')
    else:
        doc.add_heading('AST-Based Detection Results', level=1)
        doc.add_paragraph('✅ No vulnerabilities detected using AST analysis.')
    
    # Add regex-based analysis
    doc.add_heading('Pattern-Based Analysis', level=1)
    doc.add_paragraph('Code with highlighted vulnerable patterns:')
    
    para = doc.add_paragraph()
    i = 0
    while i < len(highlighted):
        if highlighted.startswith('[VULNERABLE:', i):
            end = highlighted.find(']', i)
            if end != -1:
                vuln_text = highlighted[i+12:end]
                run = para.add_run(vuln_text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                i = end + 1
                continue
        para.add_run(highlighted[i])
        i += 1
    
    # Add summary
    doc.add_heading('Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f'AST Analysis: {len(ast_vulnerabilities)} vulnerabilities found\n')
    summary.add_run(f'Pattern Analysis: {highlighted.count("[VULNERABLE:")} potential issues identified\n')
    
    if ast_vulnerabilities or highlighted.count("[VULNERABLE:") > 0:
        summary.add_run('\n⚠️  This code contains potential SQL injection vulnerabilities.')
    else:
        summary.add_run('\n✅ No SQL injection vulnerabilities detected.')
    
    doc.save(docx_filename)
    return docx_filename

def is_github_py_url(url):
    return url.startswith('https://github.com/') and url.endswith('.py')

def github_raw_url(url):
    # Convert GitHub blob URL to raw URL
    # e.g. https://github.com/user/repo/blob/main/file.py -> https://raw.githubusercontent.com/user/repo/main/file.py
    if '/blob/' in url:
        parts = url.split('/blob/')
        return parts[0].replace('github.com', 'raw.githubusercontent.com') + '/' + parts[1]
    return url

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            # Create JWT token
            token = jwt.encode({
                'username': username,
                'exp': datetime.utcnow() + timedelta(hours=1)  # Token expires in 1 hours
            }, app.config['JWT_SECRET_KEY'], algorithm='HS256')
            
            return jsonify({
                'token': token,
                'username': username,
                'message': 'Login successful'
            }), 200
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Login failed: {str(e)}'}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    return jsonify({'message': 'Logout successful'}), 200

@app.route('/api/verify-token', methods=['POST'])
@token_required
def verify_token(current_user):
    return jsonify({
        'valid': True,
        'username': current_user
    }), 200

@app.route('/api/scan', methods=['POST'])
@token_required
def api_scan(current_user):
    ensure_dirs()
    
    url = request.form.get('url', '').strip()
    code = request.form.get('code', '').strip()
    file = request.files.get('file')
    
    vulnerabilities = []
    scan_info = {
        'scan_timestamp': datetime.now().isoformat(),
        'input_type': None,
        'file_name': None
    }
    
    source_code = ""
    
    try:
        if url:
            # GitHub URL scanning
            scan_info['input_type'] = 'url'
            scan_info['file_name'] = url
            
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        python_code = response.text
                        source_code = python_code
                        
                        # Save to temporary file for AST analysis
                        temp_filename = f"temp_github_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                        temp_filepath = os.path.join('results', temp_filename)
                        with open(temp_filepath, 'w', encoding='utf-8') as f:
                            f.write(python_code)
                        
                        # Run AST-based analysis
                        vulnerabilities = ast_detector.scan_file(temp_filepath)
                        
                        # Clean up temporary file
                        if os.path.exists(temp_filepath):
                            os.remove(temp_filepath)
                    else:
                        return jsonify({'error': f"Failed to fetch GitHub file: {response.status_code}"}), 400
                except Exception as e:
                    return jsonify({'error': f"Error fetching GitHub URL: {str(e)}"}), 400
            else:
                return jsonify({'error': "Please provide a valid GitHub Python file URL"}), 400
        
        elif file and file.filename:
            # File upload scanning
            scan_info['input_type'] = 'file'
            scan_info['file_name'] = file.filename
            
            filename = file.filename
            if filename.endswith('.py'):
                try:
                    content = file.read().decode('utf-8')
                    source_code = content
                    
                    # Save to temporary file for AST analysis
                    temp_filename = f"temp_upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                    temp_filepath = os.path.join('results', temp_filename)
                    with open(temp_filepath, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    # Run AST-based analysis
                    vulnerabilities = ast_detector.scan_file(temp_filepath)
                    
                    # Clean up temporary file
                    if os.path.exists(temp_filepath):
                        os.remove(temp_filepath)
                        
                except Exception as e:
                    return jsonify({'error': f"Error processing uploaded file: {str(e)}"}), 400
            else:
                return jsonify({'error': "Please upload a Python (.py) file"}), 400
        
        elif code:
            # Code pasting scanning
            scan_info['input_type'] = 'code'
            scan_info['file_name'] = 'pasted_code.py'
            source_code = code
            
            try:
                # Save to temporary file for AST analysis
                temp_filename = f"temp_paste_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                temp_filepath = os.path.join('results', temp_filename)
                with open(temp_filepath, 'w', encoding='utf-8') as f:
                    f.write(code)
                
                # Run AST-based analysis
                vulnerabilities = ast_detector.scan_file(temp_filepath)
                
                # Clean up temporary file
                if os.path.exists(temp_filepath):
                    os.remove(temp_filepath)
                    
            except Exception as e:
                return jsonify({'error': f"Error processing pasted code: {str(e)}"}), 400
        
        else:
            return jsonify({'error': "Please provide a GitHub URL, upload a file, or paste code"}), 400
        
        # Convert vulnerabilities to JSON-serializable format
        vulnerabilities_json = []
        for vuln in vulnerabilities:
            vulnerabilities_json.append({
                'file_path': vuln.file_path,
                'line_number': vuln.line_number,
                'vulnerability_type': vuln.vulnerability_type,
                'description': vuln.description,
                'severity': vuln.severity,
                'code_snippet': vuln.code_snippet,
                'remediation': vuln.remediation,
                'confidence': vuln.confidence
            })
        
        # Calculate summary statistics
        summary = {
            'total_vulnerabilities': len(vulnerabilities_json),
            'high_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'high'),
            'medium_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'medium'),
            'low_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'low')
        }
        
        # Create highlighted source code with vulnerabilities marked in red
        highlighted_code = highlight_sql_injection_web(source_code)
        
        return jsonify({
            'vulnerabilities': vulnerabilities_json,
            'summary': summary,
            'scan_info': scan_info,
            'highlighted_code': highlighted_code,
            'original_code': source_code,
            'scan_timestamp': scan_info['scan_timestamp']
        })
        
    except Exception as e:
        return jsonify({'error': f"Unexpected error: {str(e)}"}), 500

@app.route('/api/generate-report', methods=['POST'])
@token_required
def generate_word_report(current_user):
    """Generate a Word document report from scan results"""
    try:
        data = request.get_json()
        
        # Extract data from request
        vulnerabilities = data.get('vulnerabilities', [])
        summary = data.get('summary', {})
        scan_info = data.get('scan_info', {})
        original_code = data.get('original_code', '')
        
        # Create Word document
        doc = Document()
        
        # Add title
        title = doc.add_heading('SQL Injection Security Scan Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add file information
        doc.add_heading('Scan Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        # Populate info table
        info_data = [
            ['Scan Date:', scan_info.get('scan_timestamp', 'N/A')],
            ['Input Type:', scan_info.get('input_type', 'N/A')],
            ['File Name:', scan_info.get('file_name', 'N/A')],
            ['Total Vulnerabilities:', str(summary.get('total_vulnerabilities', 0))]
        ]
        
        for i, (key, value) in enumerate(info_data):
            row = info_table.rows[i]
            row.cells[0].text = key
            row.cells[0].paragraphs[0].runs[0].bold = True
            row.cells[1].text = value
        
        # Add executive summary
        doc.add_heading('Executive Summary', level=1)
        summary_para = doc.add_paragraph()
        
        total_vulns = summary.get('total_vulnerabilities', 0)
        high_severity = summary.get('high_severity', 0)
        medium_severity = summary.get('medium_severity', 0)
        low_severity = summary.get('low_severity', 0)
        
        if total_vulns > 0:
            summary_para.add_run('⚠️ SECURITY ISSUES DETECTED\n\n').bold = True
            summary_para.add_run(f'This security scan identified {total_vulns} potential SQL injection vulnerabilities:\n')
            summary_para.add_run(f'• High Severity: {high_severity} issues\n')
            summary_para.add_run(f'• Medium Severity: {medium_severity} issues\n')
            summary_para.add_run(f'• Low Severity: {low_severity} issues\n\n')
            
            if high_severity > 0:
                summary_para.add_run('IMMEDIATE ACTION REQUIRED: High severity vulnerabilities pose significant security risks and should be addressed immediately.')
        else:
            summary_para.add_run('✅ NO VULNERABILITIES DETECTED\n\n').bold = True
            summary_para.add_run('This security scan found no SQL injection vulnerabilities in the analyzed code.')
        
        # Add vulnerability details
        if vulnerabilities:
            doc.add_heading('Vulnerability Details', level=1)
            
            for i, vuln in enumerate(vulnerabilities, 1):
                # Vulnerability header
                vuln_heading = doc.add_heading(f'Vulnerability #{i}: {vuln.get("vulnerability_type", "").replace("_", " ").title()}', level=2)
                
                # Vulnerability info table
                vuln_table = doc.add_table(rows=6, cols=2)
                vuln_table.style = 'Table Grid'
                
                vuln_data = [
                    ['File Path:', vuln.get('file_path', 'N/A')],
                    ['Line Number:', str(vuln.get('line_number', 'N/A'))],
                    ['Severity:', vuln.get('severity', 'N/A')],
                    ['Confidence:', f"{vuln.get('confidence', 0):.2f}"],
                    ['Description:', vuln.get('description', 'N/A')],
                    ['Remediation:', vuln.get('remediation', 'N/A')]
                ]
                
                for j, (key, value) in enumerate(vuln_data):
                    row = vuln_table.rows[j]
                    row.cells[0].text = key
                    row.cells[0].paragraphs[0].runs[0].bold = True
                    row.cells[1].text = value
                
                # Add code snippet with vulnerability highlighting if available
                if vuln.get('code_snippet'):
                    doc.add_paragraph('Code Snippet:')
                    
                    # Highlight vulnerable patterns in the code snippet
                    highlighted_snippet = highlight_sql_injection_word(vuln['code_snippet'])
                    
                    # Parse and format the highlighted code snippet
                    code_para = doc.add_paragraph()
                    code_para.style = 'No Spacing'
                    
                    i = 0
                    while i < len(highlighted_snippet):
                        if highlighted_snippet.startswith('[VULNERABLE:', i):
                            # Find the end of the vulnerable section
                            end = highlighted_snippet.find(']', i)
                            if end != -1:
                                # Extract the vulnerable code (remove the marker)
                                vuln_text = highlighted_snippet[i+12:end]
                                # Add the vulnerable code with red formatting
                                run = code_para.add_run(vuln_text)
                                run.font.name = 'Courier New'
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                                run.bold = True
                                i = end + 1
                                continue
                        
                        # Add normal character
                        run = code_para.add_run(highlighted_snippet[i])
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        i += 1
                
                doc.add_paragraph()  # Add spacing
        
        # Add recommendations
        doc.add_heading('Security Recommendations', level=1)
        recommendations = [
            'Use parameterized queries or prepared statements instead of string concatenation',
            'Implement input validation and sanitization',
            'Use ORM (Object-Relational Mapping) frameworks when possible',
            'Apply the principle of least privilege to database accounts',
            'Regularly update and patch database systems',
            'Implement proper error handling that doesn\'t expose sensitive information',
            'Use web application firewalls (WAF) as an additional layer of protection',
            'Conduct regular security code reviews and penetration testing'
        ]
        
        for rec in recommendations:
            doc.add_paragraph(f'• {rec}')
        
        # Add source code section with vulnerability highlighting
        if original_code and len(original_code.strip()) > 0:
            doc.add_heading('Analyzed Source Code', level=1)
            doc.add_paragraph('The following source code was analyzed for SQL injection vulnerabilities (vulnerable patterns highlighted in red):')
            
            # Generate highlighted code with vulnerability markers
            highlighted_code = highlight_sql_injection_word(original_code)
            
            # Parse and format the highlighted code
            code_para = doc.add_paragraph()
            code_para.style = 'No Spacing'
            
            i = 0
            while i < len(highlighted_code):
                if highlighted_code.startswith('[VULNERABLE:', i):
                    # Find the end of the vulnerable section
                    end = highlighted_code.find(']', i)
                    if end != -1:
                        # Extract the vulnerable code (remove the marker)
                        vuln_text = highlighted_code[i+12:end]
                        # Add the vulnerable code with red formatting
                        run = code_para.add_run(vuln_text)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(8)
                        run.font.color.rgb = RGBColor(255, 0, 0)  # Red color
                        run.bold = True
                        i = end + 1
                        continue
                
                # Add normal character
                run = code_para.add_run(highlighted_code[i])
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                i += 1
        
        # Save document to temporary file
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'security_scan_report_{timestamp}.docx'
        
        # Create temporary file for download
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        # Return file directly for download
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f"Failed to generate report: {str(e)}"}), 500

@app.route('/api/enhanced-scan', methods=['POST'])
@token_required
def enhanced_api_scan(current_user):
    """Enhanced API endpoint for scanning with SonarQube security standards"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        # Get scan parameters
        scan_type = data.get('scan_type', 'code')  # 'code', 'file', 'url'
        code_content = data.get('code', '')
        file_content = data.get('file_content', '')
        url = data.get('url', '')
        
        results = {}
        
        if scan_type == 'code' and code_content:
            # Scan provided code
            results = _scan_code_content_enhanced(code_content, 'user_input.py')
        elif scan_type == 'file' and file_content:
            # Scan uploaded file content
            results = _scan_code_content_enhanced(file_content, 'uploaded_file.py')
        elif scan_type == 'url' and url:
            # Scan GitHub URL
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        results = _scan_code_content_enhanced(response.text, url)
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'Invalid scan parameters'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during enhanced scan: {str(e)}'}), 500

@app.route('/api/sonarqube-export', methods=['POST'])
@token_required
def sonarqube_export(current_user):
    """Export vulnerabilities in SonarQube format"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'error': 'No data provided'}), 400
        
        vulnerabilities = data.get('vulnerabilities', [])
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400
        
        # Convert to SonarQube format
        sonar_issues = []
        for vuln in vulnerabilities:
            sonar_issue = {
                "engineId": "python-security-scanner",
                "ruleId": vuln.get('rule_key', 'python:S2077'),
                "severity": vuln.get('severity', 'MAJOR'),
                "type": "VULNERABILITY",
                "primaryLocation": {
                    "message": vuln.get('description', 'SQL Injection vulnerability'),
                    "filePath": vuln.get('file_path', 'unknown'),
                    "textRange": {
                        "startLine": vuln.get('line_number', 1),
                        "endLine": vuln.get('line_number', 1)
                    }
                },
                "cwe": vuln.get('cwe_references', []),
                "owasp": vuln.get('owasp_references', []),
                "confidence": vuln.get('confidence', 0.5)
            }
            sonar_issues.append(sonar_issue)
        
        # Create temporary file
        temp_file = tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json')
        json.dump({"issues": sonar_issues}, temp_file, indent=2)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sonarqube_issues_{datetime.now().strftime("%Y%m%d_%H%M%S")}.json',
            mimetype='application/json'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error exporting SonarQube format: {str(e)}'}), 500

@app.route('/api/security-standards', methods=['GET'])
@token_required
def get_security_standards(current_user):
    """Get available security standards and categories"""
    try:
        standards = {
            "sq_categories": [
                {
                    "key": category.key,
                    "name": category.name.replace('_', ' ').title(),
                    "vulnerability_probability": category.vulnerability.name
                }
                for category in SQCategory
            ],
            "vulnerability_probabilities": [
                {
                    "name": prob.name,
                    "score": prob.value
                }
                for prob in VulnerabilityProbability
            ],
            "cwe_mappings": {
                "sql_injection": ["89", "564", "943"],
                "nosql_injection": ["89", "943"],
                "command_injection": ["77", "78", "88", "214"]
            },
            "owasp_top10_2021": [
                "A01:2021-Broken Access Control",
                "A02:2021-Cryptographic Failures",
                "A03:2021-Injection",
                "A04:2021-Insecure Design",
                "A05:2021-Security Misconfiguration",
                "A06:2021-Vulnerable and Outdated Components",
                "A07:2021-Identification and Authentication Failures",
                "A08:2021-Software and Data Integrity Failures",
                "A09:2021-Security Logging and Monitoring Failures",
                "A10:2021-Server-Side Request Forgery"
            ]
        }
        
        return jsonify(standards)
        
    except Exception as e:
        return jsonify({'error': f'Error getting security standards: {str(e)}'}), 500

def _scan_code_content_enhanced(code_content: str, source_name: str) -> dict:
    """Enhanced scanning of code content with SonarQube security standards"""
    try:
        # Create temporary file for scanning
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(code_content)
        temp_file.close()
        
        # Create a fresh detector instance for each scan to avoid cached vulnerabilities
        fresh_detector = EnhancedSQLInjectionDetector()
        vulnerabilities = fresh_detector.scan_file(temp_file.name)
        
        # Clean up temporary file
        os.unlink(temp_file.name)
        
        # Generate enhanced report
        report = fresh_detector.get_enhanced_report()
        summary = report.get('summary', {})
        
        # Generate highlighted code for React frontend
        highlighted_code = None
        original_code = code_content
        file_name = source_name
        
        if vulnerabilities:
            # Use existing highlight function to highlight vulnerable patterns
            highlighted_code = highlight_sql_injection_web(code_content)
            
            # Escape HTML entities in the original code for safe display
            original_code = code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            
            # Extract filename from source
            if '/' in source_name:
                file_name = source_name.split('/')[-1]
            elif source_name.startswith('http'):
                file_name = source_name.split('/')[-1] if '/' in source_name else 'scanned_code.py'
        
        # Format results with both enhanced and UI-compatible formats
        results = {
            'source': source_name,
            'scan_type': 'enhanced',
            'summary': summary,
            'compliance': report.get('compliance', {}),
            'vulnerabilities': [vuln.to_dict() for vuln in vulnerabilities],
            'total_vulnerabilities': len(vulnerabilities),
            'scan_timestamp': datetime.now().isoformat(),
            
            # Additional UI-compatible fields
            'total_issues': len(vulnerabilities),
            'high_severity': summary.get('high_severity', 0),
            'medium_severity': summary.get('medium_severity', 0),
            'low_severity': summary.get('low_severity', 0),
            'high_count': summary.get('high', 0),
            'medium_count': summary.get('medium', 0),
            'low_count': summary.get('low', 0),
            
            # Code highlighting fields for React frontend
            'highlighted_code': highlighted_code,
            'original_code': original_code,
            'file_name': file_name
        }
        
        return results
        
    except Exception as e:
        return {
            'error': f'Error during enhanced scan: {str(e)}',
            'source': source_name,
            'scan_type': 'enhanced',
            'vulnerabilities': [],
            'total_vulnerabilities': 0,
            'total_issues': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'highlighted_code': None,
            'original_code': '',
            'file_name': source_name
        }

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    debug = os.environ.get('FLASK_ENV') != 'production'
    app.run(host='0.0.0.0', port=port, debug=debug)