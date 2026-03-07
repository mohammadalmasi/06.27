from flask import request, jsonify, send_file
import requests
import re
import os
import ast
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from datetime import datetime
import json
import tempfile

class SQLInjectionVulnerability:
    def __init__(self, line_number, vulnerability_type, description, severity, code_snippet, remediation, confidence, file_path=None):
        self.line_number = line_number
        self.vulnerability_type = vulnerability_type
        self.description = description
        self.severity = severity
        self.code_snippet = code_snippet
        self.remediation = remediation
        self.confidence = confidence
        self.file_path = file_path or 'unknown'
        
    def to_dict(self):
        return {
            'line_number': self.line_number,
            'vulnerability_type': self.vulnerability_type,
            'description': self.description,
            'severity': self.severity,
            'code_snippet': self.code_snippet,
            'remediation': self.remediation,
            'confidence': self.confidence,
            'file_path': self.file_path,
            'cwe_references': ["89", "564", "943"],
            'owasp_references': ["A03:2021-Injection"],
            'rule_key': 'python:S2077'
        }

def scan_file(filename):
    """Scan a Python file for SQL injection vulnerabilities. Returns list of SQLInjectionVulnerability."""
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            code = f.read()
    except UnicodeDecodeError:
        with open(filename, 'r', encoding='latin-1') as f:
            code = f.read()
    try:
        tree = ast.parse(code)
        analyzer = SQLInjectionTaintAnalyzer(filename, code)
        analyzer.analyze(tree)
        return analyzer.vulnerabilities
    except SyntaxError:
        return []

class SQLInjectionTaintAnalyzer:
    """
    Taint + sink analysis for SQL injection.
    Taint sources = untrusted data (input(), request.args, etc.).
    Sinks = dangerous operations (cursor.execute(query), text(sql), etc.).
    If tainted data reaches a sink -> potential vulnerability.
    """
    # Calls that introduce tainted data (sources)
    TAINT_SOURCE_ATTRS = frozenset({
        'args', 'form', 'cookies', 'headers', 'json', 'data', 'values',
        'get', 'getlist', 'get_json', 'get_data'
    })
    TAINT_SOURCE_NAMES = frozenset({'input'})
    REQUEST_LIKE_NAMES = frozenset({'request', 'req', 'flask_request', 'environ'})
    # Method names that are SQL sinks (first argument is the query string)
    SQL_SINK_ATTRS = frozenset({'execute', 'executemany', 'raw'})
    SQL_SINK_NAMES = frozenset({'text'})

    def __init__(self, filename, source_code):
        self.filename = filename
        self.source_code = source_code
        self.vulnerabilities = []
        self._assignments = []   # (line, list of target name ids, value node)
        self._sink_calls = []    # (line, call node, arg index for SQL)
        self._tainted = set()

    def analyze(self, tree):
        self._collect_assignments_and_sinks(tree)
        self._compute_tainted_fixpoint()
        self._report_tainted_sinks()

    def _collect_assignments_and_sinks(self, node):
        """Gather all Assign/AugAssign and sink calls in execution order (DFS)."""
        if isinstance(node, ast.Assign):
            targets = []
            for t in node.targets:
                if isinstance(t, ast.Name):
                    targets.append(t.id)
            if targets:
                self._assignments.append((getattr(node, 'lineno', 0), targets, node.value))
        elif isinstance(node, ast.AugAssign) and isinstance(node.target, ast.Name):
            # x += expr propagates taint to x
            self._assignments.append((getattr(node, 'lineno', 0), [node.target.id], node.value))
        elif isinstance(node, ast.Call):
            line = getattr(node, 'lineno', 0)
            if isinstance(node.func, ast.Attribute):
                attr = getattr(node.func, 'attr', None)
                if attr in self.SQL_SINK_ATTRS:
                    if node.args:
                        self._sink_calls.append((line, node, 0))
            elif isinstance(node.func, ast.Name) and getattr(node.func, 'id', None) in self.SQL_SINK_NAMES:
                if node.args:
                    self._sink_calls.append((line, node, 0))
        for child in ast.iter_child_nodes(node):
            self._collect_assignments_and_sinks(child)

    def _is_taint_source_call(self, node):
        """True if this Call is a taint source (e.g. input(), request.args.get())."""
        if not isinstance(node, ast.Call):
            return False
        if isinstance(node.func, ast.Name):
            return getattr(node.func, 'id', None) in self.TAINT_SOURCE_NAMES
        if isinstance(node.func, ast.Attribute):
            attr = getattr(node.func, 'attr', None)
            if attr not in self.TAINT_SOURCE_ATTRS:
                return False
            val = node.func.value
            if isinstance(val, ast.Name):
                return val.id in self.REQUEST_LIKE_NAMES
            if isinstance(val, ast.Attribute) and isinstance(getattr(val, 'value', None), ast.Name):
                return getattr(val.value, 'id', None) in self.REQUEST_LIKE_NAMES
            if attr in ('get', 'getlist', 'get_json', 'get_data'):
                return True
        if isinstance(node.func, ast.Subscript):
            return self._is_request_like(node.func.value)
        return False

    def _is_request_like(self, node):
        """request, request.args, request.form, etc."""
        if isinstance(node, ast.Name):
            return node.id in self.REQUEST_LIKE_NAMES
        if isinstance(node, ast.Attribute):
            if getattr(node, 'attr', None) in self.TAINT_SOURCE_ATTRS:
                return self._is_request_like(node.value)
        return False

    def _expr_tainted(self, node):
        """True if expression node uses tainted data (names, concat, f-string, source call)."""
        if node is None:
            return False
        if isinstance(node, ast.Name):
            return getattr(node, 'id', None) in self._tainted
        if isinstance(node, ast.Call):
            if self._is_taint_source_call(node):
                return True
            for a in node.args:
                if self._expr_tainted(a):
                    return True
            for k in getattr(node, 'keywords', []) or []:
                if self._expr_tainted(k.value):
                    return True
            return False
        if isinstance(node, ast.BinOp) and isinstance(getattr(node, 'op', None), ast.Add):
            return self._expr_tainted(node.left) or self._expr_tainted(node.right)
        if isinstance(node, ast.JoinedStr):
            for v in node.values:
                if isinstance(v, ast.FormattedValue) and self._expr_tainted(v.value):
                    return True
            return False
        if isinstance(node, ast.Attribute):
            return self._expr_tainted(node.value) if hasattr(node, 'value') else False
        if isinstance(node, ast.Subscript):
            # request.form['x'], request.args['id'] etc. are taint sources
            if self._is_request_like(node.value):
                return True
            return self._expr_tainted(node.value) or self._expr_tainted(
                node.slice if isinstance(node.slice, ast.AST) else None
            )
        if isinstance(node, (ast.List, ast.Tuple)):
            for e in (getattr(node, 'elts', []) or []):
                if self._expr_tainted(e):
                    return True
            return False
        return False

    def _compute_tainted_fixpoint(self):
        """Compute set of tainted variable names until fixpoint."""
        changed = True
        while changed:
            changed = False
            for _line, targets, value in self._assignments:
                if not self._expr_tainted(value):
                    continue
                for name in targets:
                    if name not in self._tainted:
                        self._tainted.add(name)
                        changed = True

    def _report_tainted_sinks(self):
        """Report vulnerability when a sink receives tainted argument."""
        for line, call_node, arg_idx in self._sink_calls:
            if arg_idx >= len(call_node.args):
                continue
            arg = call_node.args[arg_idx]
            if not self._expr_tainted(arg):
                continue
            snippet = self._get_code_snippet(call_node)
            method = 'execute' if isinstance(call_node.func, ast.Attribute) else 'text()'
            vulnerability = SQLInjectionVulnerability(
                line_number=line,
                vulnerability_type='sql_injection',
                description=f'Tainted data (user input) flows to SQL sink ({method}) - SQL injection risk',
                severity='high',
                code_snippet=snippet,
                remediation='Use parameterized queries; do not build SQL from user input.',
                confidence=0.9,
                file_path=self.filename
            )
            self.vulnerabilities.append(vulnerability)

    def _get_code_snippet(self, node):
        try:
            return ast.unparse(node)
        except AttributeError:
            return f'Line {getattr(node, "lineno", "unknown")}'

def highlight_sql_injection_vulnerabilities(code, vulnerabilities=None):
    """Return code as-is; UI should use vulnerabilities[].line_number and .severity to highlight lines."""
    return code if code else ''

def highlight_sql_injection_vulnerabilities_word(code):
    """Highlight SQL injection vulnerability patterns for Word documents"""
    patterns = [
        # Updated patterns to match our fixes
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{\}.*?[\'\"]\s*\.format\s*\([^)]*\))',
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?%[sd].*?[\'\"]\s*%\s*[^;]+)',
        r'(f[\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*?\{[^}]*\}.*?[\'"])',
        r'([\'\"]\s*.*(?:SELECT|INSERT|UPDATE|DELETE|CREATE|DROP|ALTER|EXEC|EXECUTE).*[\'\"]\s*\+\s*\w+)',
        r'(\w+\s*\+\s*[\'"].*(?:SELECT|INSERT|UPDATE|DELETE|WHERE|FROM|JOIN|UNION).*[\'"])',
        r'(cursor\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(\.execute\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(cursor\.execute\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(\.raw\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(\.raw\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(text\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(text\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'([\'"]WHERE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]LIKE\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'([\'"]ORDER\s+BY\s+[^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(request\.args\.get\([^)]*\))',
        r'(request\.form\.get\([^)]*\))',
        r'(request\.(?:form|args)\[[^]]*\])',
        # Multi-line concatenation patterns
        r'([\'\"]\s*(?:SELECT|INSERT|UPDATE|DELETE)[^\'\"]*[\'\"]\s*\+[^+]*\+[^+]*\+)',
        r'([\'\"]\s*WHERE[^\'\"]*[\'\"]\s*\+[^+]*\+)',
        r'([\'\"]\s*FROM[^\'\"]*[\'\"]\s*\+[^+]*\+)',
        # NoSQL injection patterns
        r'(collection\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(\.find\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input)[^}]*\})',
        r'(db\.eval\s*\(\s*f[\'"][^\'\"]*\{[^}]*\}[^\'\"]*[\'"])',
        r'(db\.eval\s*\(\s*[\'"][^\'\"]*[\'"]\s*\+\s*\w+)',
        r'(db\.eval\s*\([^)]*(?:request\.|username|user_input|\w+_input))',
        r'(\.(?:find_one|update|delete|remove|insert)\s*\(\s*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))',
        r'(\{[^}]*[\'"]:\s*request\.(?:form|args|json)\[[^]]*\][^}]*\})',
        r'(aggregate\s*\(\s*\[[^]]*\{[^}]*[\'"]:\s*(?:request\.|username|user_input|\w+_input))'
    ]
    
    highlighted = code
    for pattern in patterns:
        highlighted = re.sub(pattern, lambda m: f'[SQL-INJECTION-VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE | re.DOTALL)
    
    return highlighted

def scan_code_content_for_sql_injection(code_content: str, source_name: str) -> dict:
    """Scan code content for SQL injection vulnerabilities"""
    try:
        # Create temporary file for scanning
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(code_content)
        temp_file.close()
        
        vulnerabilities = scan_file(temp_file.name)
        os.unlink(temp_file.name)

        file_name = source_name
        if '/' in source_name:
            file_name = source_name.split('/')[-1]
        elif source_name.startswith('http'):
            file_name = source_name.split('/')[-1] if '/' in source_name else 'scanned_code.py'

        lines_to_highlight = [{'line_number': v.line_number, 'severity': v.severity} for v in vulnerabilities]

        # Calculate summary
        summary = {
            'total_vulnerabilities': len(vulnerabilities),
            'high_severity': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium_severity': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low_severity': sum(1 for v in vulnerabilities if v.severity == 'low'),
            'high': sum(1 for v in vulnerabilities if v.severity == 'high'),
            'medium': sum(1 for v in vulnerabilities if v.severity == 'medium'),
            'low': sum(1 for v in vulnerabilities if v.severity == 'low')
        }
        
        results = {
            'source': source_name,
            'scan_type': 'sql_injection',
            'summary': summary,
            'vulnerabilities': [vuln.to_dict() for vuln in vulnerabilities],
            'lines_to_highlight': lines_to_highlight,
            'code': code_content,
            'original_code': code_content,
            'highlighted_code': code_content.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'),
            'total_vulnerabilities': len(vulnerabilities),
            'scan_timestamp': datetime.now().isoformat(),
            'total_issues': len(vulnerabilities),
            'high_severity': summary['high_severity'],
            'medium_severity': summary['medium_severity'],
            'low_severity': summary['low_severity'],
            'high_count': summary['high'],
            'medium_count': summary['medium'],
            'low_count': summary['low'],
            'file_name': file_name
        }
        
        return results
        
    except Exception as e:
        return {
            'error': f'Error during SQL injection scan: {str(e)}',
            'source': source_name,
            'scan_type': 'sql_injection',
            'vulnerabilities': [],
            'lines_to_highlight': [],
            'code': '',
            'original_code': '',
            'highlighted_code': '',
            'total_vulnerabilities': 0,
            'total_issues': 0,
            'high_severity': 0,
            'medium_severity': 0,
            'low_severity': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'file_name': source_name
        }

def is_github_py_url(url):
    """Check if URL is a GitHub Python file URL"""
    return 'github.com' in url and url.endswith('.py')

def github_raw_url(url):
    """Convert GitHub blob URL to raw URL"""
    return url.replace('github.com', 'raw.githubusercontent.com').replace('/blob/', '/')

def api_scan_sql_injection(current_user):
    """API endpoint for SQL injection scanning"""
    try:
        data = request.get_json()
        code_content = data.get('code')
        url = data.get('url')
        
        if code_content:
            # Scan provided code content
            results = scan_code_content_for_sql_injection(code_content, 'Direct input')
            
        elif url:
            # Scan URL content
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        results = scan_code_content_for_sql_injection(response.text, url)
                    else:
                        return jsonify({'error': f'Failed to fetch URL: {response.status_code}'}), 400
                except Exception as e:
                    return jsonify({'error': f'Error fetching URL: {str(e)}'}), 400
            else:
                return jsonify({'error': 'Invalid GitHub Python file URL'}), 400
        else:
            return jsonify({'error': 'Invalid scan parameters'}), 400
        
        return jsonify(results)
        
    except Exception as e:
        return jsonify({'error': f'Error during SQL injection scan: {str(e)}'}), 500

def api_generate_sql_injection_report(current_user):
    try:
        data = request.get_json()
        vulnerabilities = data.get('vulnerabilities', [])
        source = data.get('source', 'Unknown')
        
        if not vulnerabilities:
            return jsonify({'error': 'No vulnerabilities provided'}), 400

        doc = Document()
        title = doc.add_heading('SQL Injection Security Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_heading('Report Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_table.cell(0, 0).text = 'Source'
        info_table.cell(0, 1).text = source
        info_table.cell(1, 0).text = 'Report Generated'
        info_table.cell(1, 1).text = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        info_table.cell(2, 0).text = 'Total Vulnerabilities'
        info_table.cell(2, 1).text = str(len(vulnerabilities))
        info_table.cell(3, 0).text = 'Risk Level'
        info_table.cell(3, 1).text = 'High' if len(vulnerabilities) > 0 else 'Low'
        doc.add_heading('Executive Summary', level=1)
        summary_text = f"""
This report presents the results of a comprehensive SQL injection vulnerability analysis performed on the provided code. 
The analysis identified {len(vulnerabilities)} potential SQL injection vulnerabilities that could allow attackers to manipulate 
database queries and potentially gain unauthorized access to sensitive data.

SQL injection vulnerabilities can lead to data breaches, data modification, data deletion, and in some cases, complete system compromise.
These vulnerabilities should be addressed immediately to prevent potential attacks against the application's database.
        """
        doc.add_paragraph(summary_text.strip())
        doc.add_heading('Vulnerability Details', level=1)
        
        for i, vuln in enumerate(vulnerabilities, 1):
            doc.add_heading(f'SQL Injection Vulnerability #{i}', level=2)
            vuln_table = doc.add_table(rows=5, cols=2)
            vuln_table.style = 'Table Grid'
            
            vuln_table.cell(0, 0).text = 'Line Number'
            vuln_table.cell(0, 1).text = str(vuln.get('line_number', 'N/A'))
            vuln_table.cell(1, 0).text = 'Severity'
            vuln_table.cell(1, 1).text = vuln.get('severity', 'Medium').title()
            vuln_table.cell(2, 0).text = 'CWE References'
            vuln_table.cell(2, 1).text = ', '.join(vuln.get('cwe_references', []))
            vuln_table.cell(3, 0).text = 'OWASP References'
            vuln_table.cell(3, 1).text = ', '.join(vuln.get('owasp_references', []))
            vuln_table.cell(4, 0).text = 'Description'
            vuln_table.cell(4, 1).text = vuln.get('description', 'SQL injection vulnerability detected')
            doc.add_paragraph('Vulnerable Code:', style='Heading 3')
            code_para = doc.add_paragraph()
            code_run = code_para.add_run(vuln.get('code_snippet', ''))
            code_run.font.name = 'Courier New'
            code_run.font.size = Pt(10)
            code_run.font.color.rgb = RGBColor(255, 0, 0)
            doc.add_paragraph('Remediation:', style='Heading 3')
            remediation_para = doc.add_paragraph(vuln.get('remediation', 'Use parameterized queries and proper input validation'))
            doc.add_paragraph()

        doc.add_heading('SQL Injection Prevention Recommendations', level=1)
        recommendations = """
1. Parameterized Queries: Always use parameterized queries or prepared statements instead of string concatenation.
2. Input Validation: Validate all user inputs on both client and server side.
3. Least Privilege: Use database accounts with minimal necessary privileges.
4. Stored Procedures: Use stored procedures when possible, but ensure they are also parameterized.
5. Escape Special Characters: If parameterized queries are not possible, properly escape special characters.
6. ORM Usage: Use Object-Relational Mapping (ORM) frameworks that handle parameterization automatically.
7. Database Firewalls: Implement database firewalls to detect and block SQL injection attempts.
8. Regular Security Testing: Conduct regular security testing including automated SQL injection scanning.
9. Code Reviews: Implement thorough code review processes to catch SQL injection vulnerabilities.
10. Security Training: Provide security awareness training to developers about SQL injection risks.
        """
        doc.add_paragraph(recommendations.strip())
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name=f'sql_injection_security_report_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': f'Error generating SQL injection report: {str(e)}'}), 500

def test_sql_injection_scanner():
    """Test function to verify the SQL injection scanner is working"""
    test_cases = [
        # SQLAlchemy text() cases
        'query = text(f"SELECT * FROM users WHERE id = {user_id}")',
        'result = text(f"SELECT * FROM products WHERE category = {category}")',
        'sql = text("SELECT * FROM users WHERE id = " + str(user_id))',
        
        # F-string cases
        'sql = f"SELECT * FROM users WHERE name = {username}"',
        'query = f"INSERT INTO logs VALUES ({user_id}, {message})"',
        
        # String concatenation cases
        'query = "SELECT * FROM users WHERE id = " + user_id',
        'sql = "DELETE FROM users WHERE name = " + username',
        
        # cursor.execute cases
        'cursor.execute(f"SELECT * FROM users WHERE id = {user_id}")',
        'cursor.execute("SELECT * FROM users WHERE id = " + str(user_id))',
        
        # Django ORM raw cases
        'User.objects.raw(f"SELECT * FROM users WHERE id = {user_id}")',
        'results = MyModel.objects.raw("SELECT * FROM table WHERE id = " + user_id)',
        
        # NoSQL injection cases
        'collection.find({"name": request.form["username"]})',
        'db.eval(f"this.name == {username}")',
        
        # Safe cases (should NOT be detected)
        'query = text("SELECT * FROM users WHERE id = :user_id", user_id=user_id)',
        'cursor.execute("SELECT * FROM users WHERE id = %s", (user_id,))',
        'safe_query = "SELECT * FROM users"',
    ]
    
    print("Testing SQL Injection Scanner...")
    print("=" * 50)
    
    for i, test_case in enumerate(test_cases, 1):
        print(f"\nTest {i}: {test_case}")
        
        # Create temporary file with test case
        temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.py', delete=False)
        temp_file.write(test_case)
        temp_file.close()
        
        # Scan the test case
        vulnerabilities = scan_file(temp_file.name)
        os.unlink(temp_file.name)
        
        # Report results
        if vulnerabilities:
            print(f"  ✓ DETECTED: {len(vulnerabilities)} vulnerability(ies)")
            for vuln in vulnerabilities:
                print(f"    - {vuln.description}")
        else:
            # Check if this is a safe case that should NOT be detected
            if any(safe_pattern in test_case for safe_pattern in [':user_id', '%s', '(user_id,)', 'SELECT * FROM users"']):
                print(f"  ✓ SAFE: No vulnerabilities detected (expected)")
            else:
                print(f"  ✗ MISSED: No vulnerabilities detected (unexpected)")
    
    print("\n" + "=" * 50)
    print("Test completed. If you see any 'MISSED' results, there may be an issue with the scanner.")
    
# Test the scanner if run directly
if __name__ == "__main__":
    test_sql_injection_scanner() 


# cd /Users/mohammen.almasi/thesis/06.27 && source venv/bin/activate && cd backend && python main.py

# cd thesis/06.27/frontend && npm start