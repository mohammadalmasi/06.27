from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import re
import os
from docx import Document
from docx.shared import RGBColor
import sqlite3
from sql_injection_detector import SQLInjectionDetector
from datetime import datetime, timedelta
import jwt
from functools import wraps

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 2 * 1024 * 1024  # 2MB upload limit
app.config['JWT_SECRET_KEY'] = 'sql-injection-scanner-secret-key-2024'  # Change this in production

# Enable CORS for React frontend
CORS(app, origins=["http://localhost:3000", "http://127.0.0.1:3000"], 
     methods=["GET", "POST", "OPTIONS"], 
     allow_headers=["Content-Type", "Authorization"])

# Initialize the advanced SQL injection detector
ast_detector = SQLInjectionDetector()

# Authentication configuration
ADMIN_USERNAME = "Admin"
ADMIN_PASSWORD = "unipassau"

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = None
        
        # JWT is passed in the request header
        if 'Authorization' in request.headers:
            auth_header = request.headers['Authorization']
            try:
                token = auth_header.split(" ")[1]  # Bearer <token>
            except IndexError:
                return jsonify({'error': 'Token is missing!'}), 401
        
        if not token:
            return jsonify({'error': 'Token is missing!'}), 401
        
        try:
            # Decode the token
            data = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
            current_user = data['username']
        except jwt.ExpiredSignatureError:
            return jsonify({'error': 'Token has expired!'}), 401
        except jwt.InvalidTokenError:
            return jsonify({'error': 'Token is invalid!'}), 401
        
        return f(current_user, *args, **kwargs)
    
    return decorated

INDEX_HTML = '''
<!DOCTYPE html>
<html>
<head>
    <title>SQL Injection Web Scanner</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; }
        .container { max-width: 600px; margin: auto; }
        input[type=text], input[type=file], textarea { width: 80%; padding: 8px; }
        textarea { height: 120px; }
        input[type=submit] { padding: 8px 16px; }
    </style>
</head>
<body>
    <div class="container">
        <h2>SQL Injection Web Scanner</h2>
        <form method="post" action="/scan" enctype="multipart/form-data">
            <label for="url">Enter GitHub Python file URL (or leave blank to upload or paste):</label><br>
            <input type="text" id="url" name="url" placeholder="https://github.com/user/repo/blob/main/file.py"><br><br>
            <label for="file">Or upload a Python source file:</label><br>
            <input type="file" id="file" name="file"><br><br>
            <label for="code">Or paste Python source code:</label><br>
            <textarea id="code" name="code" placeholder="# Paste your Python code here..."></textarea><br><br>
            <input type="submit" value="Scan">
        </form>
    </div>
</body>
</html>
'''

RESULTS_HTML = '''
<!DOCTYPE html>
<html>
<head>
    <title>Scan Results</title>
    <style>
        body { font-family: Arial, sans-serif; margin: 40px; background-color: #f5f5f5; }
        .container { max-width: 1200px; margin: auto; background-color: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
        .vuln { color: #d32f2f; font-weight: bold; background-color: #ffebee; padding: 2px 4px; border-radius: 3px; }
        pre { background: #f8f9fa; padding: 15px; border-radius: 6px; border-left: 4px solid #007bff; overflow-x: auto; font-size: 14px; line-height: 1.4; }
        h2 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }
        h3 { color: #34495e; margin-top: 25px; margin-bottom: 15px; }
        h4 { color: #e74c3c; margin-top: 20px; margin-bottom: 10px; }
        .download-link { display: inline-block; background-color: #28a745; color: white; padding: 10px 20px; text-decoration: none; border-radius: 5px; margin: 10px 0; }
        .download-link:hover { background-color: #218838; }
        .file-summary { background-color: #e9ecef; padding: 15px; border-radius: 5px; margin: 20px 0; }
        .error { color: #dc3545; background-color: #f8d7da; padding: 10px; border-radius: 5px; border: 1px solid #f5c6cb; }
    </style>
</head>
<body>
    <div class="container">
        <h2>Scan Results for {{ url or filename or 'Uploaded File' }}</h2>
        {% if results %}
            {% for page, code in results.items() %}
                <h3>{{ page }}</h3>
                {% if page == 'Word Document' %}
                    {{ code|safe }}
                {% elif page == 'File Summary' %}
                    <div class="file-summary">{{ code|safe }}</div>
                {% elif page == 'Error' %}
                    <div class="error">{{ code|safe }}</div>
                {% else %}
                    <pre>{{ code|safe }}</pre>
                {% endif %}
            {% endfor %}
        {% else %}
            <p>No vulnerabilities found or no code detected.</p>
        {% endif %}
        <a href="/" style="display: inline-block; margin-top: 20px; padding: 10px 20px; background-color: #007bff; color: white; text-decoration: none; border-radius: 5px;">&#8592; New Scan</a>
    </div>
</body>
</html>
'''

def crawl_site(start_url, max_pages=10):
    """Crawl the site starting from start_url, return {url: html} for each page."""
    visited = set()
    to_visit = [start_url]
    results = {}
    domain = urlparse(start_url).netloc
    
    while to_visit and len(visited) < max_pages:
        url = to_visit.pop(0)
        if url in visited:
            continue
        try:
            resp = requests.get(url, timeout=5)
            if resp.status_code != 200:
                continue
            html = resp.text
            results[url] = html
            visited.add(url)
            soup = BeautifulSoup(html, 'html.parser')
            for link in soup.find_all('a', href=True):
                next_url = urljoin(url, link['href'])
                next_domain = urlparse(next_url).netloc
                if next_domain == domain and next_url not in visited and next_url not in to_visit:
                    to_visit.append(next_url)
        except Exception as e:
            continue
    return results

def extract_code_blocks(html):
    """Extract <script> tags and form actions from HTML."""
    soup = BeautifulSoup(html, 'html.parser')
    code_blocks = []
    # Extract <script> tags
    for script in soup.find_all('script'):
        if script.string:
            code_blocks.append(('JavaScript', script.string.strip()))
    # Extract form actions and inline event handlers
    for form in soup.find_all('form'):
        if form.get('action'):
            code_blocks.append(('Form action', form['action']))
        # Inline event handlers
        for attr in form.attrs:
            if attr.startswith('on') and form[attr]:
                code_blocks.append((f'Form {attr}', form[attr]))
    # Extract inline event handlers from all elements
    for tag in soup.find_all(True):
        for attr in tag.attrs:
            if attr.startswith('on') and tag[attr]:
                code_blocks.append((f'{tag.name} {attr}', tag[attr]))
    return code_blocks

def highlight_sql_injection(code):
    """Highlight vulnerable SQL and NoSQL patterns in code (JS/Python-like)."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def highlight_sql_injection_web(code):
    """Highlight vulnerable SQL and NoSQL patterns in code for web display."""
    # More comprehensive patterns for SQL and NoSQL injection detection
    patterns = [
        # SQL injection patterns
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        r'\{\w+\}',     # {user_input} in f-strings
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
        # NoSQL injection patterns (MongoDB)
        r'(find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find({"name": username})
        r'(find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # find_one({"name": username})
        r'(db\.[\w_]+\.find\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find({"name": username})
        r'(db\.[\w_]+\.find_one\s*\(\s*\{[^}]*[\w\'\"]+\s*:\s*\w+[^}]*\}\s*\))',  # db.users.find_one({"name": username})
    ]
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'<span class="vuln">{m.group(0)}</span>', highlighted, flags=re.IGNORECASE)
    return highlighted

def ensure_dirs():
    os.makedirs('sourcecodes', exist_ok=True)
    os.makedirs('results', exist_ok=True)

def save_code_block(page_idx, code_type, code):
    ensure_dirs()
    ext = 'js' if 'js' in code_type.lower() or 'script' in code_type.lower() else 'txt'
    filename = f"sourcecodes/page{page_idx}_{code_type.replace(' ', '_').lower()}.{ext}"
    with open(filename, 'w', encoding='utf-8') as f:
        f.write(code)
    return filename

def result_txt_to_docx(result_filename):
    if not result_filename.endswith('.py') and not result_filename.endswith('.txt'):
        return
    docx_filename = result_filename.rsplit('.', 1)[0] + '.docx'
    with open(result_filename, 'r', encoding='utf-8') as f:
        code = f.read()
    doc = Document()
    para = doc.add_paragraph()
    i = 0
    while i < len(code):
        if code.startswith('[VULNERABLE:', i):
            end = code.find(']', i)
            if end != -1:
                vuln_text = code[i+12:end]
                run = para.add_run(vuln_text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                i = end + 1
                continue
        para.add_run(code[i])
        i += 1
    doc.save(docx_filename)
    return docx_filename

def scan_code_file_ast(filename):
    """Scan code file using AST-based SQL injection detector"""
    try:
        vulnerabilities = ast_detector.scan_file(filename)
        return vulnerabilities
    except Exception as e:
        print(f"AST detection error: {e}")
        return []

def scan_code_file(filename):
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            code = f.read()
    except UnicodeDecodeError:
        with open(filename, 'r', encoding='latin-1') as f:
            code = f.read()
    
    # Get AST-based vulnerabilities
    ast_vulnerabilities = scan_code_file_ast(filename)
    
    # More comprehensive patterns for SQL injection detection (regex-based)
    patterns = [
        # String concatenation with user input
        r'(\+\s*\w+)',  # + user_input
        r'(\+\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # "SELECT * FROM users WHERE id = " + user_input
        
        # F-string formatting
        r'\{\w+\}',     # {user_input} in f-strings
        
        # String formatting
        r'%\s*\w+',      # % user_input
        r'\.format\(\w+\)',  # .format(user_input)
        
        # Database query patterns with user input
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\+\s*\w+)',  # execute("SELECT * FROM " + table)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*%\s*\w+)',   # execute("SELECT * FROM %s" % user_input)
        r'(execute\s*\(\s*[\'"][^\'"]*[\'"]\s*\.\s*format\s*\(\s*\w+)',  # execute("SELECT * FROM {}".format(user_input))
        
        # Raw SQL queries with variables
        r'(SELECT\s+.*\s+FROM\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',  # SELECT * FROM users WHERE id = + user_input
        r'(INSERT\s+INTO\s+.*\s+VALUES\s*\(\s*.*\s*\+\s*\w+)',  # INSERT INTO users VALUES ( + user_input)
        r'(UPDATE\s+.*\s+SET\s+.*\s+WHERE\s+.*\s*\+\s*\w+)',   # UPDATE users SET name = + user_input
        
        # Common vulnerable patterns
        r'(request\.form\[\w+\])',  # request.form['user_input']
        r'(request\.args\[\w+\])',  # request.args['user_input']
        r'(request\.cookies\[\w+\])',  # request.cookies['user_input']
        r'(input\s*\(\s*[\'"][^\'"]*[\'"]\s*\))',  # input("Enter value: ")
    ]
    
    highlighted = code
    for pat in patterns:
        highlighted = re.sub(pat, lambda m: f'[VULNERABLE:{m.group(0)}]', highlighted, flags=re.IGNORECASE)
    
    # Create comprehensive report with both detection methods
    base = os.path.basename(filename)
    docx_filename = os.path.join('results', base.replace('.', '_result.', 1).rsplit('.', 1)[0] + '.docx')
    doc = Document()
    
    # Add title
    title = doc.add_heading('SQL Injection Analysis Report', 0)
    
    # Add file information
    doc.add_heading('File Information', level=1)
    doc.add_paragraph(f'File: {filename}')
    doc.add_paragraph(f'Analysis Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add AST-based vulnerabilities
    if ast_vulnerabilities:
        doc.add_heading('AST-Based Detection Results', level=1)
        doc.add_paragraph(f'Found {len(ast_vulnerabilities)} vulnerabilities using AST analysis:')
        
        for vuln in ast_vulnerabilities:
            # Vulnerability details
            p = doc.add_paragraph()
            p.add_run(f'Line {vuln.line_number}: ').bold = True
            p.add_run(f'{vuln.description}')
            
            # Severity and confidence
            p2 = doc.add_paragraph()
            p2.add_run(f'Severity: {vuln.severity} | Confidence: {vuln.confidence:.2f}')
            
            # Code snippet
            if vuln.code_snippet:
                doc.add_paragraph('Code Snippet:')
                code_para = doc.add_paragraph(vuln.code_snippet)
                code_para.style = 'No Spacing'
            
            # Remediation
            doc.add_paragraph(f'Remediation: {vuln.remediation}')
            doc.add_paragraph('---')
    else:
        doc.add_heading('AST-Based Detection Results', level=1)
        doc.add_paragraph('✅ No vulnerabilities detected using AST analysis.')
    
    # Add regex-based analysis
    doc.add_heading('Pattern-Based Analysis', level=1)
    doc.add_paragraph('Code with highlighted vulnerable patterns:')
    
    para = doc.add_paragraph()
    i = 0
    while i < len(highlighted):
        if highlighted.startswith('[VULNERABLE:', i):
            end = highlighted.find(']', i)
            if end != -1:
                vuln_text = highlighted[i+12:end]
                run = para.add_run(vuln_text)
                run.font.color.rgb = RGBColor(255, 0, 0)
                run.bold = True
                i = end + 1
                continue
        para.add_run(highlighted[i])
        i += 1
    
    # Add summary
    doc.add_heading('Summary', level=1)
    summary = doc.add_paragraph()
    summary.add_run(f'AST Analysis: {len(ast_vulnerabilities)} vulnerabilities found\n')
    summary.add_run(f'Pattern Analysis: {highlighted.count("[VULNERABLE:")} potential issues identified\n')
    
    if ast_vulnerabilities or highlighted.count("[VULNERABLE:") > 0:
        summary.add_run('\n⚠️  This code contains potential SQL injection vulnerabilities.')
    else:
        summary.add_run('\n✅ No SQL injection vulnerabilities detected.')
    
    doc.save(docx_filename)
    return docx_filename

def is_github_py_url(url):
    return url.startswith('https://github.com/') and url.endswith('.py')

def github_raw_url(url):
    # Convert GitHub blob URL to raw URL
    # e.g. https://github.com/user/repo/blob/main/file.py -> https://raw.githubusercontent.com/user/repo/main/file.py
    if '/blob/' in url:
        parts = url.split('/blob/')
        return parts[0].replace('github.com', 'raw.githubusercontent.com') + '/' + parts[1]
    return url

@app.route('/api/login', methods=['POST'])
def login():
    try:
        data = request.get_json()
        username = data.get('username')
        password = data.get('password')
        
        if username == ADMIN_USERNAME and password == ADMIN_PASSWORD:
            # Create JWT token
            token = jwt.encode({
                'username': username,
                'exp': datetime.utcnow() + timedelta(hours=24)  # Token expires in 24 hours
            }, app.config['JWT_SECRET_KEY'], algorithm='HS256')
            
            return jsonify({
                'token': token,
                'username': username,
                'message': 'Login successful'
            }), 200
        else:
            return jsonify({'error': 'Invalid credentials'}), 401
            
    except Exception as e:
        return jsonify({'error': f'Login failed: {str(e)}'}), 500

@app.route('/api/logout', methods=['POST'])
def logout():
    return jsonify({'message': 'Logout successful'}), 200

@app.route('/api/verify-token', methods=['POST'])
@token_required
def verify_token(current_user):
    return jsonify({
        'valid': True,
        'username': current_user
    }), 200

@app.route('/api/scan', methods=['POST'])
@token_required
def api_scan(current_user):
    ensure_dirs()
    
    url = request.form.get('url', '').strip()
    code = request.form.get('code', '').strip()
    file = request.files.get('file')
    
    vulnerabilities = []
    scan_info = {
        'scan_timestamp': datetime.now().isoformat(),
        'input_type': None,
        'file_name': None
    }
    
    source_code = ""
    
    try:
        if url:
            # GitHub URL scanning
            scan_info['input_type'] = 'url'
            scan_info['file_name'] = url
            
            if is_github_py_url(url):
                raw_url = github_raw_url(url)
                try:
                    response = requests.get(raw_url, timeout=10)
                    if response.status_code == 200:
                        python_code = response.text
                        source_code = python_code
                        
                        # Save to temporary file for AST analysis
                        temp_filename = f"temp_github_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                        temp_filepath = os.path.join('results', temp_filename)
                        with open(temp_filepath, 'w', encoding='utf-8') as f:
                            f.write(python_code)
                        
                        # Run AST-based analysis
                        vulnerabilities = ast_detector.scan_file(temp_filepath)
                        
                        # Clean up temporary file
                        if os.path.exists(temp_filepath):
                            os.remove(temp_filepath)
                    else:
                        return jsonify({'error': f"Failed to fetch GitHub file: {response.status_code}"}), 400
                except Exception as e:
                    return jsonify({'error': f"Error fetching GitHub URL: {str(e)}"}), 400
            else:
                return jsonify({'error': "Please provide a valid GitHub Python file URL"}), 400
        
        elif file and file.filename:
            # File upload scanning
            scan_info['input_type'] = 'file'
            scan_info['file_name'] = file.filename
            
            filename = file.filename
            if filename.endswith('.py'):
                try:
                    content = file.read().decode('utf-8')
                    source_code = content
                    
                    # Save to temporary file for AST analysis
                    temp_filename = f"temp_upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                    temp_filepath = os.path.join('results', temp_filename)
                    with open(temp_filepath, 'w', encoding='utf-8') as f:
                        f.write(content)
                    
                    # Run AST-based analysis
                    vulnerabilities = ast_detector.scan_file(temp_filepath)
                    
                    # Clean up temporary file
                    if os.path.exists(temp_filepath):
                        os.remove(temp_filepath)
                        
                except Exception as e:
                    return jsonify({'error': f"Error processing uploaded file: {str(e)}"}), 400
            else:
                return jsonify({'error': "Please upload a Python (.py) file"}), 400
        
        elif code:
            # Code pasting scanning
            scan_info['input_type'] = 'code'
            scan_info['file_name'] = 'pasted_code.py'
            source_code = code
            
            try:
                # Save to temporary file for AST analysis
                temp_filename = f"temp_paste_{datetime.now().strftime('%Y%m%d_%H%M%S')}.py"
                temp_filepath = os.path.join('results', temp_filename)
                with open(temp_filepath, 'w', encoding='utf-8') as f:
                    f.write(code)
                
                # Run AST-based analysis
                vulnerabilities = ast_detector.scan_file(temp_filepath)
                
                # Clean up temporary file
                if os.path.exists(temp_filepath):
                    os.remove(temp_filepath)
                    
            except Exception as e:
                return jsonify({'error': f"Error processing pasted code: {str(e)}"}), 400
        
        else:
            return jsonify({'error': "Please provide a GitHub URL, upload a file, or paste code"}), 400
        
        # Convert vulnerabilities to JSON-serializable format
        vulnerabilities_json = []
        for vuln in vulnerabilities:
            vulnerabilities_json.append({
                'file_path': vuln.file_path,
                'line_number': vuln.line_number,
                'vulnerability_type': vuln.vulnerability_type,
                'description': vuln.description,
                'severity': vuln.severity,
                'code_snippet': vuln.code_snippet,
                'remediation': vuln.remediation,
                'confidence': vuln.confidence
            })
        
        # Calculate summary statistics
        summary = {
            'total_vulnerabilities': len(vulnerabilities_json),
            'high_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'high'),
            'medium_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'medium'),
            'low_severity': sum(1 for v in vulnerabilities_json if v['severity'].lower() == 'low')
        }
        
        # Create highlighted source code with vulnerabilities marked in red
        highlighted_code = highlight_sql_injection_web(source_code)
        
        return jsonify({
            'vulnerabilities': vulnerabilities_json,
            'summary': summary,
            'scan_info': scan_info,
            'highlighted_code': highlighted_code,
            'original_code': source_code
        })
        
    except Exception as e:
        return jsonify({'error': f"Unexpected error: {str(e)}"}), 500

@app.route('/download/<filename>')
def download_file(filename):
    """Route to download generated Word documents"""
    try:
        file_path = os.path.join('results', filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return "File not found", 404
    except Exception as e:
        return f"Error downloading file: {e}", 500

@app.route('/user')
def get_user():
    user_id = request.args.get('id')  # <-- Source
    conn = sqlite3.connect('db.sqlite')
    cursor = conn.cursor()
    query = "SELECT * FROM users WHERE id = " + user_id  # <-- Sink
    cursor.execute(query)
    return cursor.fetchall()

if __name__ == '__main__':
    app.run(debug=True, port=5001) 
    # http://127.0.0.1:5001/